"""文档读取模块 - 支持Word和PDF文件读取"""
import logging
import subprocess
from pathlib import Path
from typing import Optional
from docx import Document
import pdfplumber

logger = logging.getLogger(__name__)


class DocumentReader:
    """文档读取器"""
    
    # 支持的Word文件扩展名
    WORD_EXTENSIONS = {'.docx', '.doc'}
    # 支持的PDF文件扩展名
    PDF_EXTENSIONS = {'.pdf'}
    
    @staticmethod
    def check_dependencies():
        """
        检查系统依赖（如antiword）是否已安装
        如果未安装，直接抛出异常并退出程序
        """
        try:
            result = subprocess.run(
                ['antiword', '-h'],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=5
            )
            # antiword 返回非0也可能表示已安装（某些版本）
            # 只要命令能执行（不抛出FileNotFoundError）就认为已安装
        except FileNotFoundError:
            error_msg = (
                "错误: antiword 未安装。\n\n"
                "请先安装 antiword 工具：\n"
                "  - Mac: brew install antiword\n"
                "  - Linux: sudo apt-get install antiword 或 sudo yum install antiword\n"
                "  - Windows: 下载并安装antiword，或使用chocolatey: choco install antiword\n\n"
                "安装完成后请重新运行程序。"
            )
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except Exception as e:
            # 其他错误（如超时）也认为可能有问题，但先继续运行
            logger.warning(f"检查antiword时出现警告: {e}")
    
    @staticmethod
    def scan_documents(data_dir: Path) -> list[Path]:
        """
        扫描目录，获取所有Word和PDF文件
        
        Args:
            data_dir: 数据目录路径
            
        Returns:
            文件路径列表（按创建时间倒序排列，最新的文件在前）
        """
        files = []
        
        if not data_dir.exists():
            logger.warning(f"数据目录不存在: {data_dir}")
            return files
        
        for ext in DocumentReader.WORD_EXTENSIONS | DocumentReader.PDF_EXTENSIONS:
            files.extend(data_dir.glob(f"*{ext}"))
        
        # 去重并按创建时间倒序排序（最新的文件在前）
        files = list(set(files))
        
        # 按照文件创建时间倒序排序
        def get_creation_time(file_path: Path) -> float:
            """获取文件创建时间"""
            try:
                stat = file_path.stat()
                # macOS 使用 st_birthtime，Linux/Windows 使用 st_ctime
                if hasattr(stat, 'st_birthtime'):
                    return stat.st_birthtime
                else:
                    return stat.st_ctime
            except (OSError, AttributeError):
                return 0.0
        
        files.sort(key=get_creation_time, reverse=True)
        logger.info(f"扫描到 {len(files)} 个文档文件（按创建时间倒序排列，最新的文件优先处理）")
        
        return files
    
    @staticmethod
    def read_word(file_path: Path) -> Optional[str]:
        """
        读取Word文档内容（支持.docx和.doc格式）
        - .docx文件使用textract库
        - .doc文件使用antiword工具
        
        Args:
            file_path: Word文件路径
            
        Returns:
            文档文本内容，失败抛出异常
            
        Raises:
            ImportError: textract库未安装（.docx文件）
            FileNotFoundError: antiword未安装（.doc文件）
            RuntimeError: 文件读取失败
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return DocumentReader._read_docx(file_path)
        elif suffix == '.doc':
            return DocumentReader._read_doc(file_path)
        else:
            logger.warning(f"不支持的Word文件格式: {suffix}")
            return None
    
    @staticmethod
    def _read_docx(file_path: Path) -> Optional[str]:
        """
        读取.docx文件（使用python-docx库）
        
        Args:
            file_path: .docx文件路径
            
        Returns:
            文档文本内容，失败返回None
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            content = "\n".join(paragraphs)
            
            # 处理表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    content += "\n" + "\n".join(table_text)
            
            if content:
                logger.info(f"使用python-docx成功读取.docx文件: {file_path.name}")
                return content
            else:
                logger.warning(f"python-docx读取.docx文件返回空内容: {file_path.name}")
                return None
                
        except Exception as e:
            logger.error(f"读取.docx文件失败 {file_path.name}: {e}")
            return None
    
    @staticmethod
    def _read_doc(file_path: Path) -> Optional[str]:
        """
        读取.doc文件（使用antiword工具）
        
        Args:
            file_path: .doc文件路径
            
        Returns:
            文档文本内容，失败抛出异常
            
        Raises:
            FileNotFoundError: antiword未安装
            RuntimeError: antiword执行失败
        """
        try:
            # 调用antiword命令提取文本
            # 使用绝对路径避免特殊字符问题
            abs_path = file_path.resolve()
            result = subprocess.run(
                ['antiword', str(abs_path)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                encoding='utf-8',
                errors='ignore',
                timeout=30,
                cwd=str(file_path.parent)  # 设置工作目录
            )
            
            if result.returncode == 0:
                content = result.stdout.strip()
                if content:
                    logger.info(f"使用antiword成功读取.doc文件: {file_path.name}")
                    return content
                else:
                    logger.warning(f"antiword读取.doc文件返回空内容: {file_path.name}")
                    return None
            else:
                error_msg = f"无法读取.doc文件 {file_path.name}：antiword执行失败。错误信息: {result.stderr.strip()}"
                logger.error(error_msg)
                raise RuntimeError(error_msg)
                
        except FileNotFoundError:
            # 这种情况理论上不应该发生，因为启动时已经检查过了
            error_msg = f"无法读取.doc文件 {file_path.name}：antiword未安装。请确保已安装antiword工具。"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        except subprocess.TimeoutExpired:
            error_msg = f"读取.doc文件超时: {file_path.name}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except Exception as e:
            error_msg = f"读取.doc文件时出错 {file_path.name}：{str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg) from e
    
    @staticmethod
    def read_pdf(file_path: Path) -> Optional[str]:
        """
        读取PDF文档内容
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            文档文本内容，失败返回None
        """
        try:
            content_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text.strip())
            
            content = "\n".join(content_parts)
            return content if content else None
            
        except Exception as e:
            logger.error(f"读取PDF文件失败 {file_path}: {e}")
            return None
    
    @staticmethod
    def read_document(file_path: Path) -> Optional[str]:
        """
        读取文档内容（自动识别文件类型）
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            文档文本内容，失败返回None
        """
        suffix = file_path.suffix.lower()
        
        if suffix in DocumentReader.WORD_EXTENSIONS:
            return DocumentReader.read_word(file_path)
        elif suffix in DocumentReader.PDF_EXTENSIONS:
            return DocumentReader.read_pdf(file_path)
        else:
            logger.warning(f"不支持的文件类型: {suffix}")
            return None
    
    @staticmethod
    def preprocess_text(text: str, max_length: int = 50000) -> str:
        """
        文本预处理
        
        Args:
            text: 原始文本
            max_length: 最大长度限制（避免超出LLM token限制）
            
        Returns:
            处理后的文本
        """
        if not text:
            return ""
        
        # 去除多余空白字符
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 限制长度
        if len(text) > max_length:
            text = text[:max_length] + "\n[文本已截断...]"
            logger.warning(f"文本长度超过限制，已截断至 {max_length} 字符")
        
        return text
